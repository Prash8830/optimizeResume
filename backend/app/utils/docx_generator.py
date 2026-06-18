import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

SECTION_HEADERS = {"SUMMARY", "EXPERIENCE", "PROJECTS", "TECHNICAL SKILLS", "SKILLS", "EDUCATION"}


def generate_docx(resume_text: str, role_title: str = "Resume") -> bytes:
    doc = Document()
    _set_default_font(doc)

    sections = _parse_sections(resume_text)

    # ── HEADER ──────────────────────────────────────────────────────────────
    header_lines = sections.pop("HEADER", [])
    if header_lines:
        name_line = next((l.strip() for l in header_lines if l.strip()), "")
        contact_line = next((l.strip() for l in header_lines[1:] if l.strip()), "")
        if name_line:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(name_line)
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        if contact_line:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(contact_line)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    else:
        p = doc.add_paragraph()
        run = p.add_run(role_title)
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    # ── BODY SECTIONS ───────────────────────────────────────────────────────
    for section_name, lines in sections.items():
        if section_name == "OVERVIEW":
            continue
        heading = doc.add_heading(section_name.upper(), level=1)
        _style_heading(heading)

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue
            para = doc.add_paragraph()
            run = para.add_run(stripped)
            run.font.size = Pt(9.5)
            # Bold role/company lines in EXPERIENCE
            if section_name == "EXPERIENCE" and " at " in stripped and not stripped[0].islower():
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)


def _style_heading(heading) -> None:
    for run in heading.runs:
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)


def _parse_sections(text: str) -> dict[str, list[str]]:
    ALL_HEADERS = SECTION_HEADERS | {"HEADER"}
    sections: dict[str, list[str]] = {}
    current = "OVERVIEW"
    lines: list[str] = []

    for line in text.split("\n"):
        upper = line.strip().upper()
        if upper in ALL_HEADERS:
            if lines or current != "OVERVIEW":
                sections[current] = lines
            current = upper
            lines = []
        else:
            lines.append(line)

    if lines:
        sections[current] = lines
    return sections
